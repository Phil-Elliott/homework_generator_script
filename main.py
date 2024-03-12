from docx import Document
import os
from docx.shared import Pt

def main(lesson_number, vocabulary, grammar, practice_questions):
    doc = Document()
    
    # Class Review Materials Heading
    doc.add_heading(f'Class Review Materials {lesson_number}', level=0)

    # Vocabulary Section
    doc.add_heading('Vocabulary', level=1)
    for word in vocabulary:
        doc.add_paragraph(word, style='ListBullet')
    
    # Grammar Section
    doc.add_heading('Grammar', level=1)
    for rule in grammar:
        doc.add_paragraph(rule, style='ListBullet')
    
    # Practice Questions and Answers
    letters = ['A', 'B', 'C', 'D']

    doc.add_heading('Practice Questions', level=1)
    for i in range(len(practice_questions)):
        p = doc.add_paragraph()
        p.add_run(f'Q{i+1} {practice_questions[i]["question"]}').bold = True

        letters = ['A', 'B', 'C', 'D']
        for j in range(len(practice_questions[i]["choices"])):
            p = doc.add_paragraph(f'{letters[j]}. {practice_questions[i]["choices"][j]}')
            p.paragraph_format.left_indent = Pt(20)

    doc.add_page_break()  
    doc.add_heading('Answer Key', level=1)
    for i in range(len(practice_questions)):
        doc.add_paragraph(f'Q{i+1}. {letters[i]} {practice_questions[i]["answer"]}')

    # Ensure the 'hw' folder exists
    if not os.path.exists('hw'):
        os.makedirs('hw')
    
    # Save the document
    file_path = os.path.join('hw', f'Class_Review_Materials_Lesson_{lesson_number}.docx')
    doc.save(file_path)
    
    return file_path

if __name__ == "__main__":
    lesson_number = 1 
    vocabulary = ['socks', 'shoes', 'underwear', 'this', 'that']
    grammar = ['What is this? This is a _______.  (We use this when it is close)',
               'What is that? That is a ________. (We use that when it is far)']
    practice_questions = [
        {
            "question": "What you wear on your feet?",
            "choices": [
                "I wear a shirt on my feet.",
                "I wear a jacket on my feet.",
                "I wear shoes on my feet.",
                "I wear a hat on my feet."
            ],
            "answer": "I wear shoes on my feet."
        },
        {
            "question": "What do you wear on your head?",
            "choices": [
                "I wear a shirt on my head.",
                "I wear a jacket on my head.",
                "I wear a hat on my head.",
                "I wear shoes on my head."
            ],
            "answer": "I wear a hat on my head."
        },
        {
            "question": "What do you wear on your body?",
            "choices": [
                "I wear a shirt on my body.",
                "I wear a jacket on my body.",
                "I wear a hat on my body.",
                "I wear shoes on my body."
            ],
            "answer": "I wear a shirt on my body."
        }
    ]

    
    file_path = main(lesson_number, vocabulary, grammar, practice_questions)
    print(f"Document saved to {file_path}")



# add letters to answers
# add inputs from the terminal
    
# lesson_number = input("Enter the lesson number: ")
#     vocabulary = input("Enter the vocabulary words separated by a comma: ").split(',')
#     grammar = input("Enter the grammar rules separated by a comma: ").split(',')
#     practice_questions = []
#     while True:
#         question = input("Enter the question: ")
#         choices = input("Enter the choices separated by a comma: ").split(',')
#         answer = input("Enter the answer: ")
#         practice_questions.append({
#             "question": question,
#             "choices": choices,
#             "answer": answer
#         })
#         add_more = input("Do you want to add more questions? (yes/no): ")
#         if add_more.lower() == 'no':
#             break